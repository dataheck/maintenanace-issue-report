import argparse
import json
import os
from pathlib import Path
from string import Template

import chromedriver_autoinstaller
import docx
from docx import Document
from dotenv import load_dotenv
from github import Consts
from gql import Client, gql
from gql.transport.requests import RequestsHTTPTransport
from selenium import webdriver
from selenium.webdriver.common.by import By
import selenium.webdriver.support.ui as ui
from selenium.webdriver.remote.webdriver import WebDriver
from time import sleep


PRINT_DIALOG_DELAY = 3000 # milliseconds
PRINT_SAVE_DELAY = 500    # milliseconds
INTERACTIVE_TIMEOUT = 320 # seconds


def process_configuration(env_path: str) -> dict:
    load_dotenv(dotenv_path=env_path)

    mandatory_configuration = {
        'GITHUB_API_KEY', 'GITHUB_ORGANIZATION', 'GITHUB_PROJECT_NUMBER', 'GITHUB_PROJECT_FINISHED_COLUMN',
        'PDF_SAVE_PATH', 'COVERPAGE_TEMPLATE_PATH', 'CLIENT_NAME', 'CLIENT_CONTACT', 'PROJECT_NAME', 'OUTPUT_PATH'
    }
    config = dict()

    for key in mandatory_configuration:
        value = os.environ.get(key)
        assert value is not None, f"Please set {key} before proceeding."
        config[key] = value
    
    return config


def process_graphql_project_items(result: dict , config: dict) -> list:
    completed_issues = list()

    for node in result['organization']['projectV2']['items']['nodes']:
        this_column = None
        for project_node in node['content']['projectItems']['nodes']:
            if project_node['fieldValueByName'] is not None:
                this_column = project_node['fieldValueByName']['name']
        
        assert this_column is not None, "Could not determine column for item"

        if this_column == config['GITHUB_PROJECT_FINISHED_COLUMN']:
            completed_issues.append({
                'url': node['content']['url'],
                'title': node['content']['title'],
                'number': node['content']['number'],
                'closedAt': node['content']['closedAt']
            })
    
    completed_issues = sorted(completed_issues, key=lambda x: x['closedAt'], reverse=True)

    return completed_issues

# For "projectsV2"
# assumes organization rather than user project
def initialize_github_obtain_project_column_graphql(config: dict) -> list:
    transport = RequestsHTTPTransport(
        url=Consts.DEFAULT_BASE_URL + "/graphql",
        headers={"Authorization": f"Bearer {config['GITHUB_API_KEY']}"},
    )
    
    client = Client(transport=transport, fetch_schema_from_transport=True)

    # you can find the project number by going to the project page and looking at the URL
    project_items_query = Template("""
        query {
            organization(login: \"$organization\") {
                projectV2(number: $project_number) {
                    items(first: 100, orderBy: {field: POSITION, direction: DESC }) {
                        nodes {
                            content {
                                ... on Issue {
                                    title
                                    url
                                    closedAt
                                    number
                                    projectItems(first: 100) {
                                        nodes {
                                            fieldValueByName(name: "Status") {
                                                ... on ProjectV2ItemFieldSingleSelectValue {
                                                name
                                                }
                                            }
                                        }
                                    }
                                }
                                ... on PullRequest {
                                    title
                                    url
                                    projectItems(first: 100) {
                                        nodes {
                                            fieldValueByName(name: "Status") {
                                                ... on ProjectV2ItemFieldSingleSelectValue {
                                                    name
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        }
                    }
                }
            }
        }
    """)

    query = gql(
        project_items_query.substitute(
            organization=config['GITHUB_ORGANIZATION'], 
            project_number=config['GITHUB_PROJECT_NUMBER']
        )
    )

    result : dict = client.execute(query)
    closed_issues = process_graphql_project_items(result, config)

    return closed_issues


def initialize_driver(config: dict) -> WebDriver:
    # https://stackoverflow.com/questions/47007720/set-selenium-chromedriver-userpreferences-to-save-as-pdf
    settings = {
        "recentDestinations": [{
            "id": "Save as PDF",
            "origin": "local",
            "account": "",
        }],
        "selectedDestinationId": "Save as PDF",
        "version": 2
    }

    selection_rules = {
        "kind": "local",
        "idPattern": "*",
        "namePattern": "Save as PDF"
    }

    prefs = {
        'printing.print_preview_sticky_settings.appState': json.dumps(settings), 
        'printing.default_destination_selection_rules': json.dumps(selection_rules),
        'savefile.default_directory': config['PDF_SAVE_PATH']
    }

    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_experimental_option('prefs', prefs)
    chrome_options.add_argument('--kiosk-printing')

    this_driver = webdriver.Chrome(options=chrome_options)
    
    return this_driver


class LoginTagHasValue(object):
    """ Validates that there is a meta tag with the name 'user-login' that also has non-zero-length content. """
    def __init__(self, meta_name='user-login'):
        self.meta_name = meta_name
    
    def __call__(self, this_driver):
        element = this_driver.find_element(By.XPATH, f"//meta[@name='{self.meta_name}']")
        if len(element.get_attribute("content")) > 0:
            return element
        else:
            return False


def wait_user_login(driver: WebDriver):
    driver.get("https://github.com/login")
    print("Please login to GitHub in the open window.")

    wait = ui.WebDriverWait(driver, timeout=INTERACTIVE_TIMEOUT)
    wait.until(LoginTagHasValue())


def fetch_all_issues(driver:WebDriver, config:dict, enable_print=True) -> list:
    raw_issues = initialize_github_obtain_project_column_graphql(config)
    issues = list()

    for issue in raw_issues:
        issues.append((issue['title'], issue['number'], issue['url']))
        
        if enable_print:
            driver.get(issue['url'])
            sleep(1) # otherwise we tend to miss the print dialog, sometimes.
            driver.execute_script("window.print();")
            sleep(0.5) # otherwise we tend to miss the print dialog, sometimes.

    return issues


# https://github.com/python-openxml/python-docx/issues/610#issuecomment-458289054
def add_hyperlink_into_run(paragraph, run, url):
    runs = paragraph.runs
    for i in range(len(runs)):
        if runs[i].text == run.text:
            break
    # --- This gets access to the document.xml.rels file and gets a new relation id value ---
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    # --- Create the w:hyperlink tag and add needed values ---
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    hyperlink.append(run._r)
    paragraph._p.insert(i,hyperlink)
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)


def add_issues_to_template(issues: list, config:dict) -> None:
    ''' Given a list of issues and a configuration dictionary, updates a docx template.

    The docx template is assumed to have a "List Paragraph" and "Closing Paragraph" style.
    '''
    document = Document(config['COVERPAGE_TEMPLATE_PATH'])

    document.custom_properties['ClientName'] = config['CLIENT_NAME']
    document.custom_properties['ClientContact'] = config['CLIENT_CONTACT']
    document.custom_properties['ProjectName'] = config['PROJECT_NAME']

    for title, number, url in issues:
        p = document.add_paragraph(style="List Paragraph")
        r = p.add_run()
        r.add_text(f"#{number}")
        add_hyperlink_into_run(p, r, url)
        r = p.add_run()
        r.add_text(f" - {title}")

    document.add_paragraph("If you have any questions about the above, please do not hesitate to reach out.", style="Closing Paragraph")
    document.add_paragraph("Thank you for your business.", style="Closing Paragraph")

    document.save(config['OUTPUT_PATH'])
    print("Don't forget to update fields before exporting!")
    print("When joining the PDFs, be sure to do it in chronological order (based on the modified time of the exported files). This will ensure that the table of contents matches up with the order of attached pages.")
    print("You will have to use an external tool to join the PDFs.")


def main():
    chromedriver_autoinstaller.install()

    parser = argparse.ArgumentParser(description="Collect GitHub project issues into a Word template.")
    parser.add_argument("--enable-print", help="Enable printing all issues to separate PDF files", action="store_true")
    parser.add_argument("--env", help="Path to .env file", default=".env")
    args = parser.parse_args()

    print_enabled = args.enable_print

    config = process_configuration(env_path=args.env)

    if print_enabled:
        target_path = Path(config['PDF_SAVE_PATH'])

        if not os.path.exists(target_path) and os.path.exists(target_path.parent):
            print("Creating output directory.")
            os.mkdir(target_path)
        
    assert os.path.exists(target_path), "Output location does not exist, nor does its parent - please create it or change the setting."

    if print_enabled:
        driver = initialize_driver(config)
        wait_user_login(driver)
        issues = fetch_all_issues(driver, config, enable_print=print_enabled)
        driver.close()
    else:
        issues = fetch_all_issues(None, config, enable_print=False)

    add_issues_to_template(issues, config)

if __name__ == '__main__':
    main()
