from docx import Document
from dotenv import load_dotenv
from github import Github, ProjectColumn
from pathlib import Path
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
import chromedriver_binary # pylint: disable=unused-import
import docx
import json
import os
import selenium.webdriver.support.ui as ui

PRINT_DIALOG_DELAY = 3000 # milliseconds
PRINT_SAVE_DELAY = 500    # milliseconds
INTERACTIVE_TIMEOUT = 320 # seconds

def process_configuration() -> dict:
    load_dotenv()

    mandatory_configuration = {
        'GITHUB_API_KEY', 'GITHUB_ORGANIZATION','GITHUB_PROJECT_NAME', 'GITHUB_PROJECT_FINISHED_COLUMN', 
        'PDF_SAVE_PATH', 'COVERPAGE_TEMPLATE_PATH', 'CLIENT_NAME', 'CLIENT_CONTACT', 'PROJECT_NAME', 'OUTPUT_PATH'
    }
    config = dict()

    for key in mandatory_configuration:
        value = os.environ.get(key)
        assert value is not None, f"Please set {key} before proceeding."
        config[key] = value
    
    return config


def initalize_github_obtain_project_column(config: dict) -> ProjectColumn:
    g = Github(config['GITHUB_API_KEY'])
    org = g.get_organization(config['GITHUB_ORGANIZATION'])
    project = None

    for this_project in org.get_projects():
        if this_project.name == config['GITHUB_PROJECT_NAME']:
            project = this_project
            break

    assert project is not None, f"Project '{config['GITHUB_PROJECT_NAME']}' was not found in the organization, please verify configuration."

    project_column = None

    for this_column in project.get_columns():
        if this_column.name == config['GITHUB_PROJECT_FINISHED_COLUMN']:
            project_column = this_column
            break

    assert project_column is not None, f"Project column '{config['GITHUB_PROJECT_FINISHED_COLUMN']}' was not found in the project, please verify configuration."

    return project_column


def initialize_driver(config: dict) -> WebDriver:
    # https://stackoverflow.com/questions/47007720/set-selenium-chromedriver-userpreferences-to-save-as-pdf
    settings = {
        "recentDestinations": [{
            "id": "Save as PDF",
            "origin": "local",
            "account": "",
        }],
        "selectedDestinationId": "Save as PDF",
        "version": 2
    }

    selection_rules = {
        "kind": "local",
        "idPattern": "*",
        "namePattern": "Save as PDF"
    }

    prefs = {
        'printing.print_preview_sticky_settings.appState': json.dumps(settings), 
        'printing.default_destination_selection_rules': json.dumps(selection_rules),
        'savefile.default_directory': config['PDF_SAVE_PATH']
    }

    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_experimental_option('prefs', prefs)
    chrome_options.add_argument('--kiosk-printing')

    driver = webdriver.Chrome(options=chrome_options)
    
    return driver


class LoginTagHasValue(object):
    """ Validates that there is a meta tag with the name 'user-login' that also has non-zero-length content. """
    def __init__(self, meta_name='user-login'):
        self.meta_name = meta_name
    
    def __call__(self, driver):
        element = driver.find_element(By.XPATH, f"//meta[@name='{self.meta_name}']")
        if len(element.get_attribute("content")) > 0:
            return element
        else:
            return False


def wait_user_login(driver: WebDriver):
    driver.get("https://github.com/login")
    print("Please login to GitHub in the open window.")

    wait = ui.WebDriverWait(driver, timeout=INTERACTIVE_TIMEOUT)
    wait.until(LoginTagHasValue())


def fetch_all_issues(driver:WebDriver, project_column: ProjectColumn, print=True) -> list:
    issues = list()

    for card in project_column.get_cards():
        this_issue = card.get_content()
        issues.append((this_issue.title, this_issue.number, this_issue.html_url))
        
        if print:
            driver.get(this_issue.html_url)
            driver.execute_script("window.print();")

    return issues

# https://github.com/python-openxml/python-docx/issues/610#issuecomment-458289054
def add_hyperlink_into_run(paragraph, run, url):
    runs = paragraph.runs
    for i in range(len(runs)):
        if runs[i].text == run.text:
            break
    # --- This gets access to the document.xml.rels file and gets a new relation id value ---
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    # --- Create the w:hyperlink tag and add needed values ---
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    hyperlink.append(run._r)
    paragraph._p.insert(i,hyperlink)
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)


def add_issues_to_template(issues: list, config:dict) -> None:
    ''' Given a list of issues and a configuration dictionary, updates a docx template.

    The docx template is assumed to have a "List Paragraph" and "Closing Paragraph" style.
    '''
    document = Document(config['COVERPAGE_TEMPLATE_PATH'])

    document.custom_properties['ClientName'] = config['CLIENT_NAME']
    document.custom_properties['ClientContact'] = config['CLIENT_CONTACT']
    document.custom_properties['ProjectName'] = config['PROJECT_NAME']

    for title, number, url in issues:
        p = document.add_paragraph(style="List Paragraph")
        r = p.add_run()
        r.add_text(f"#{number}")
        add_hyperlink_into_run(p, r, url)
        r = p.add_run()
        r.add_text(f" - {title}")

    document.add_paragraph("If you have any questions about the above, please do not hesitate to reach out.", style="Closing Paragraph")
    document.add_paragraph("Thank you for your business.", style="Closing Paragraph")

    document.save(config['OUTPUT_PATH'])
    print("Don't forget to update fields before exporting!")


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description="Collect GitHub project issues into a Word template.")
    parser.add_argument("--enable-print", help="Enable printing all issues to separate PDF files", action="store_true")
    args = parser.parse_args()

    print_enabled = args.enable_print

    config = process_configuration()

    if print_enabled:
        target_path = Path(config['PDF_SAVE_PATH'])

        if not os.path.exists(target_path) and os.path.exists(target_path.parent):
            print("Creating output directory.")
            os.mkdir(target_path)
    
    assert os.path.exists(config['PDF_SAVE_PATH']), "Output location does not exist, nor does its parent - please create it or change the setting."

    project_column = initalize_github_obtain_project_column(config)
    if print_enabled:
        driver = initialize_driver(config)
        wait_user_login(driver)
        issues = fetch_all_issues(driver, project_column, print=print_enabled)
        driver.close()
    else:
        issues = fetch_all_issues(None, project_column, print=False)

    add_issues_to_template(issues, config)